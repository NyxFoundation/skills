"""Notion補足資料（補助額別実施内容・工程・人月・積算根拠）をdocx化する。

NotionのPDFエクスポートで表がはみ出すため、A4横・全表フル幅のdocxを生成する。
入力: notion-fetchの保存JSON（enhanced markdown） / 図はローカル再生成済みPNGを埋め込み。
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

SOURCE_JSON = Path(sys.argv[1])
OUTPUT = Path(sys.argv[2])
FIG_DIR = Path("/home/gohan/issue52_supplement/figures")
FIGURES = {
    "図1": FIG_DIR / "fig01_project_summary.png",
    "図2": FIG_DIR / "fig04_participant_research.png",
    "図3-1": FIG_DIR / "fig03a_gantt_core.png",
    "図3-2": FIG_DIR / "fig03b_gantt_additional.png",
    "図4": FIG_DIR / "fig03_effort_matrix.png",
}
TITLE = "補足資料　補助額別実施内容・工程・人月・積算根拠"
SUBTITLE = "デジタル技術を活用した井波地域文化資源の魅力発信及び販路開拓実証事業"
CONTENT_WIDTH_MM = 269  # A4横297mm − 余白14mm×2


def load_markdown() -> str:
    data = json.loads(SOURCE_JSON.read_text())
    text = data["text"]
    return text[text.index("<content>") + len("<content>") : text.index("</content>")]


def set_run(run, size, bold=False, font="ＭＳ 明朝"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def add_inline(paragraph, text, size, bold=False):
    text = text.replace("<br>", "\n")
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1（\2）", text)
    bold_toggle = bold
    for i, chunk in enumerate(text.split("**")):
        if i > 0:
            bold_toggle = not bold_toggle if i % 2 == 1 else bold
        if not chunk:
            continue
        for j, line in enumerate(chunk.split("\n")):
            if j > 0:
                paragraph.add_run().add_break()
            if line:
                set_run(paragraph.add_run(line), size, bold_toggle)
    # 単純化: **の対応が崩れた場合も本文は欠落しない


def add_paragraph(document, text, size=10.5, bold=False, before=0, after=2, indent=None):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    add_inline(p, text, size, bold)
    return p


def table_borders(table, color="808080"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tbl_pr.append(borders)


TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
]


def table_full_width(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for ex in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(ex)
    w = OxmlElement("w:tblW")
    w.set(qn("w:type"), "pct")
    w.set(qn("w:w"), "5000")
    tbl_pr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    children = list(tbl_pr)
    order = lambda el: TBLPR_ORDER.index(el.tag.split("}")[-1]) if el.tag.split("}")[-1] in TBLPR_ORDER else len(TBLPR_ORDER)
    for child in children:
        tbl_pr.remove(child)
    for child in sorted(children, key=order):
        tbl_pr.append(child)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def cell_shading(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell(cell, text, size, bold=False):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size, bold)


def column_percents(rows):
    ncols = max(len(r) for r in rows)
    header = rows[0] if rows else []
    if ncols == 5 and any("作業" in h for h in header):
        return [8, 30, 26, 17, 19]
    if ncols == 2:
        return [26, 74]
    if ncols == 9:
        return [11, 11, 11, 11, 11, 11, 11, 11, 12]
    return [round(100 / ncols, 2)] * ncols


def add_table(document, rows):
    ncols = max(len(r) for r in rows)
    size = 10
    table = document.add_table(rows=len(rows), cols=ncols)
    table_borders(table)
    table_full_width(table)
    repeat_header(table.rows[0])
    percents = column_percents(rows)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            value = row[ci] if ci < len(row) else ""
            set_cell(cell, value.strip("*") if ri == 0 else value, size, ri == 0)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:type"), "pct")
            tc_w.set(qn("w:w"), str(int(percents[ci] * 50)))
            cell._tc.get_or_add_tcPr().append(tc_w)
            if ri == 0:
                cell_shading(cell, "EFEFEF")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(document, lines):
    table = document.add_table(rows=1, cols=1)
    table_borders(table, "B0B0B0")
    table_full_width(table)
    cell = table.cell(0, 0)
    cell_shading(cell, "F5F3EC")
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    for line in lines:
        text = line.lstrip("\t").rstrip()
        if not text:
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        if text.startswith("- "):
            p.paragraph_format.left_indent = Pt(9)
            add_inline(p, "・" + text[2:], 10)
        else:
            add_inline(p, text, 10)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document, caption):
    key = next((k for k in sorted(FIGURES, key=len, reverse=True) if caption.startswith(k)), None)
    path = FIGURES.get(key)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    if path and path.exists():
        p.add_run().add_picture(str(path), width=Mm(CONTENT_WIDTH_MM - 4))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    set_run(cap.add_run(caption), 10, True)
    return path


def parse_table_block(lines, index):
    rows = []
    current = None
    while index < len(lines):
        line = lines[index].strip()
        index += 1
        if line == "</table>":
            break
        if line == "<tr>":
            current = []
        elif line == "</tr>":
            if current:
                rows.append(current)
            current = None
        elif line.startswith("<td>"):
            current.append(re.sub(r"^<td>|</td>$", "", line))
    return rows, index


def main():
    markdown = load_markdown()
    lines = markdown.splitlines()

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Mm(14))
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Mm(13))

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run(TITLE), 16, True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_run(subtitle.add_run(SUBTITLE), 12, True)
    for p in (title, subtitle):
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    # 目次フィールド（Wordで開いた時にupdateFields設定で自動生成される）
    toc_p = document.add_paragraph()
    toc_p.paragraph_format.space_after = Pt(6)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r1 = toc_p.add_run(); r1._r.append(fld_begin)
    r2 = toc_p.add_run(); r2._r.append(instr)
    r3 = toc_p.add_run(); r3._r.append(fld_sep)
    r4 = toc_p.add_run("（目次：Wordで開くと自動生成されます）"); set_run(r4, 10)
    r5 = toc_p.add_run(); r5._r.append(fld_end)
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    figures_used = []
    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped or stripped == "---" or stripped.startswith("<colgroup") or stripped.startswith("<col ") or stripped in ("<content>", "</content>", "<table_of_contents/>"):
            index += 1
            continue
        if stripped.startswith("<callout"):
            index += 1
            inner = []
            while index < len(lines) and lines[index].strip() != "</callout>":
                inner.append(lines[index])
                index += 1
            index += 1
            add_callout(document, inner)
            continue
        if stripped.startswith("<table"):
            rows, index = parse_table_block(lines, index + 1)
            if rows:
                add_table(document, rows)
            continue
        image = re.match(r"^!\[([^\]]*)\]\(", stripped)
        if image:
            used = add_figure(document, image.group(1))
            figures_used.append((image.group(1)[:2], bool(used)))
            index += 1
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 3)
            sizes = {1: 14, 2: 12, 3: 11}
            p = document.add_paragraph(style=f"Heading {level}")
            p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
            p.paragraph_format.space_after = Pt(4)
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), str(level - 1))
            p._p.get_or_add_pPr().append(outline)
            add_inline(p, heading.group(2), sizes[level], bold=True)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = "ＭＳ ゴシック"
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "ＭＳ ゴシック")
            index += 1
            continue
        if stripped.startswith("- "):
            add_paragraph(document, "・" + stripped[2:], indent=9)
            index += 1
            continue
        add_paragraph(document, stripped)
        index += 1

    document.save(OUTPUT)

    # 検証
    check = Document(OUTPUT)
    texts = [p.text for p in check.paragraphs]
    for table in check.tables:
        for row in table.rows:
            texts.extend(c.text for c in row.cells)
    joined = "\n".join(texts)
    leftovers = [tag for tag in ("<td>", "<tr>", "<callout", "<table", "</", "<col") if tag in joined]
    required = [
        "第1版",
        "補助対象経費の下限に合わせて設定",
        "1億円=22.0人月×1,200＋直接経費7,356＝作業費33,756千円",
        "予備・単価変動枠",
        "作業番号－連番。例：C3-03",
        "6,997.5",
        "1米ドル＝160円",
        "OpenAI GPT-5.6モデル価格",
    ]
    missing = [t for t in required if t not in joined]
    report = {
        "output": str(OUTPUT),
        "bytes": OUTPUT.stat().st_size,
        "tables": len(check.tables),
        "images": len(check.inline_shapes),
        "figures": figures_used,
        "tag_leftovers": leftovers,
        "missing": missing,
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if leftovers or missing or len(check.inline_shapes) != 5:
        raise SystemExit("validation failed")


if __name__ == "__main__":
    main()
