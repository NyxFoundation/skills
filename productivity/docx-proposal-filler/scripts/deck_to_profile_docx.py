"""会社概要資料（法人紹介・別冊）をdocxで生成する。

company-deck各ページのPNGに、スピーカーノート・HP文言をもとにした補足文を添える。
様式M-3②「紹介資料・概要資料（別冊可）」への対応。
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

IMG_DIR = Path("/tmp/claude-1000/deck-final")
OUTPUT = Path("/home/gohan/nanto-original-layout-20260731/06_法人紹介資料_NyxFoundation会社紹介_20260731.docx")

PAGES = [
    ("表紙", "本資料は、一般社団法人Nyx Foundationの紹介資料です。私たちが取り組む問題、中心となる命題、これまでの実績、研究領域とプロジェクト、研究拠点、メンバー、そして今後の展望の順にご紹介します。"),
    ("中心命題 — 信頼は、検証から生まれる",
     "AIによって、文章もコードも研究計画も、もっともらしく生成できるようになりました。だからこそ問われるのは、誰が言ったかではなく、何が検証できるかです。Nyx Foundationは、形式検証、暗号、AI、セキュリティ、経済制度設計の領域を横断しながら、「検証できる信頼」を実際の成果としてつくる、東京・本郷を拠点とする独立した研究組織です。"),
    ("問題意識 — 生成コストと検証コストのギャップ",
     "AIにより文章・コード・研究計画・分析・監査レポートの生成コストは急速に下がる一方、その正しさを確かめる検証コストは、生成物が増えるほど上がっていきます。検証が追いつかなくなると、社会は信頼の根拠を中身ではなく肩書き（誰が言ったか、どの大学か、どの企業か）へ戻そうとします。しかし、肩書きはコードを検証せず、プロトコルを守らず、将来のリスクを見つけません。信頼できる理由をどうつくるかが、私たちの出発点です。"),
    ("検証できる信頼のつくり方",
     "たとえばAIがロボットを設計する場合、設計物に「どんな歩き方をしても転ばない」ことを数式で保証する数学的証明πを必ず付属させ、その証明を機械が自動で検証します。検証は誰が実行しても同じ結果が再現されます。設計者が誰であるかに関係なく、数式そのものが性質を保証する——これが、肩書きに依存しない「検証できる信頼」のつくり方です。"),
    ("実績",
     "この命題は、すでに現実の成果として動いています。①耐量子署名スキームの一部を形式検証し、ケンブリッジ大学で報告しました。②100人超が参加するグローバル監査コンペティションで、脆弱性報告件数世界1位を獲得しました。③約2.65億ドル規模のハッキングにつながりうる脆弱性を、実際の被害発生の半年前に検知・報告していました。"),
    ("6つの研究領域",
     "形式検証（仕様・実装・プロトコルを数学的に確かめる）、暗号（安全性を証明可能な形にする）、サイバーセキュリティ（リスクを事件の前に見つける）、プロトコルセキュリティ（プロトコルを信頼できる社会基盤へ）、経済制度設計（市場と制度を検証可能な対象として扱う）、公共インフラ（公共サービスに検証できる信頼を埋め込む）の6領域を横断して研究しています。"),
    ("4つのプロジェクト",
     "研究は4つのプロジェクトとして実装されています。Verity（zkVM・耐量子署名等を備えた次世代Ethereumクライアントの研究開発）、SPECA（仕様駆動のAIエージェントによる実装脆弱性の事前検知）、Eris（自律的なAIトレーダーによるDeFiプロトコルの動的検証）、経済制度分析（透明な市場データによる金融市場設計・規制議論への接続）です。実装の最前線には、本物の資金・制度・攻撃者が日々動く公共の実験場であるEthereumを選んでいます。"),
    ("共同研究・連携",
     "研究は多様な主体との連携で成り立っています。アライドアーキテクツとのEris共同研究、イーサリアム財団によるSPECA研究助成、富山県南砺市との共同記者会見、SMBC日興証券との連携、日本ブロックチェーン協会への参画、耐量子相互運用性プロジェクト、TLDR Fellowship、効率的なzkVM研究など、大学・財団・企業・行政・コミュニティと横断的につながっています。"),
    ("研究拠点 — Uzumaki",
     "強みは理念そのものではなく、理念を成果に変える装置を持っていることです。東京大学本郷キャンパスから徒歩約1分のリサーチハウス「Uzumaki」には、鋭い問いを持つ研究者が数週間から数年単位で滞在し、議論し、実装し、互いの甘さをごまかせない距離で考え続けます。散らばった着想を共有された問題へ、共有された問題を公開された成果へ変えるための拠点です。"),
    ("メンバー",
     "共同創業者のMasato KambaとIppei Torigoeを中心に、京都大学特定准教授のAkiyoshi Sannai、形式検証・耐量子暗号・LLMを横断するBanri Yanahama、東京大学経済学部のKoshi Ota、Hiroshi Tei、ZK・耐量子のShouki Tsuda、クロスチェーンのTomoki Adachiらが参画しています。肩書きではなく問いの鋭さで集まり、それぞれの問いを検証できる成果へ変えていく研究者・エンジニアの集団です。"),
    ("広がる適用領域",
     "検証できる信頼が問われる場面は、金融、行政、デジタル、教育、採用、製造、保険、物流へと広がり始めています。AIによる意思決定が増えるほど、「誰が言ったか」ではなく「何が検証できるか」が、それぞれの領域で問われていきます。"),
    ("むすび",
     "検証できる信頼は、ひとつの製品の話ではなく、社会がこれから何を信じるのかという問いです。寄付、スポンサー、共同研究など、さまざまな形でのご一緒の仕方があります。詳細は nyx.foundation をご覧ください。"),
]

PROFILE = [
    ("名称", "一般社団法人 Nyx Foundation"),
    ("所在地", "〒113-0033 東京都文京区本郷6-26-10 ニューハウジング202号室"),
    ("代表", "代表理事　鳥越 一平"),
    ("事業内容", "人工知能、数学、暗号、形式検証及び経済制度設計に係る研究開発。研究成果は論文・ソフトウェア・検証記録として公開"),
    ("研究拠点", "リサーチハウス「Uzumaki」（東京・本郷）"),
    ("ウェブサイト", "https://nyx.foundation"),
]


def set_run(run, size, bold=False, font="ＭＳ 明朝"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def main():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    for attr in ("left_margin", "right_margin"):
        setattr(section, attr, Mm(16))
    for attr in ("top_margin", "bottom_margin"):
        setattr(section, attr, Mm(12))

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run("法人紹介資料"), 20, True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("一般社団法人 Nyx Foundation　会社紹介"), 13, True)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(10)
    set_run(note.add_run("デジタル技術を活用した井波地域文化資源の魅力発信及び販路開拓実証事業　別添（2026年7月31日）"), 10.5)

    table = document.add_table(rows=len(PROFILE), cols=2)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    tbl_pr.insert(0, tbl_w)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "808080")
        borders.append(e)
    tbl_pr.append(borders)
    for (label, value), row in zip(PROFILE, table.rows):
        for cell, text, bold, pct in ((row.cells[0], label, True, 18), (row.cells[1], value, False, 82)):
            tc = cell._tc
            for child in list(tc):
                if child.tag != qn("w:tcPr"):
                    tc.remove(child)
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            set_run(p.add_run(text), 10.5, bold)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:type"), "pct")
            tc_w.set(qn("w:w"), str(pct * 50))
            cell._tc.get_or_add_tcPr().append(tc_w)

    for i, (heading, body) in enumerate(PAGES):
        h = document.add_paragraph(style="Heading 1")
        h.paragraph_format.page_break_before = True
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(f"{i + 1}. {heading}")
        set_run(run, 13, True, "ＭＳ ゴシック")
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        h._p.get_or_add_pPr().append(outline)

        img_p = document.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_after = Pt(4)
        img_p.add_run().add_picture(str(IMG_DIR / f"{i}.png"), width=Mm(240))

        body_p = document.add_paragraph()
        body_p.paragraph_format.space_after = Pt(0)
        set_run(body_p.add_run(body), 10.5)

    document.save(OUTPUT)

    check = Document(OUTPUT)
    heads = [p.text for p in check.paragraphs if p.style.name == "Heading 1"]
    sizes = {r.font.size.pt for p in check.paragraphs for r in p.runs if r.font.size}
    print("images:", len(check.inline_shapes), "/ H1:", len(heads), "/ min font:", min(sizes))
    assert len(check.inline_shapes) == 12 and len(heads) == 12 and min(sizes) >= 10
    print("OK:", OUTPUT)


if __name__ == "__main__":
    main()
