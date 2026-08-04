"""南砺市様式docx生成（2026-07-31版）。

issue #52 コメント5138640138（修正版提案書全文）を原本テンプレへ流し込む。
fill_nanto_original_layout.py がベース。変更点:
- 取得元コメントと章マーカーを現行構造（# 様式第1号／# 事業提案計画書／# 概算事業費調書）へ更新
- 必要最低補助額 5,000千円＋設定根拠の注記
- 添付書類セルへ別添3点（補足資料・法人紹介資料・実績資料）を明記
- セル内ネスト表を親セル実幅のdxa固定幅＋明示的tblGridで設定（前回の様式1幅不足の反省。pct指定は様式1本文セルにtcWが無く潰れる）
- 概算事業費調書の費目内訳を実計算値ベース（補足資料第6章と一致）へ更新
"""

import json
import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SOURCE_DIR = Path("/home/gohan/nanto-template-downloads")
OUTPUT_DIR = Path("/home/gohan/nanto-original-layout-20260731")
COMMENT_ENDPOINT = "repos/grandchildrice/life/issues/comments/5138640138"
PROJECT_NAME = "デジタル技術を活用した井波地域文化資源の魅力発信及び販路開拓実証事業"

FILES = {
    "01": (
        SOURCE_DIR / "01_事業提案書様式_記入済.docx",
        OUTPUT_DIR / "01_事業提案書様式_原本レイアウト改訂版_20260731.docx",
    ),
    "02": (
        SOURCE_DIR / "02_事業提案計画書_記入済.docx",
        OUTPUT_DIR / "02_事業提案計画書_原本レイアウト改訂版_20260731.docx",
    ),
    "03": (
        SOURCE_DIR / "03_概算事業費調書_記入済.docx",
        OUTPUT_DIR / "03_概算事業費調書_原本レイアウト改訂版_20260731.docx",
    ),
    "04": (
        SOURCE_DIR / "04_市税等納付状況確認同意書_記入済.docx",
        OUTPUT_DIR / "04_市税等納付状況確認同意書_原本レイアウト改訂版_20260731.docx",
    ),
}

MINIMUM_NOTE = (
    "概算事業費100,000千円は全ての追加工程を含む最大事業規模であり、"
    "必要最低補助額5,000千円は、既存の電子商取引環境等を利用して、対象確認、素材制作、"
    "海外向け仮説設定、少額実測、測定及び報告までを縮小して一巡できる最低額である。"
    "必要最低補助額は、企業版ふるさと納税活用事業の補助対象経費の下限（5,000千円）に合わせて設定している。"
)

ATTACHMENT_LINES = [
    "☑事業提案計画書",
    "☑概算事業費調書",
    "☑市税等納付状況確認同意書",
    "☑登記事項証明書及び定款の写し",
    "別添：補足資料『補助額別実施内容・工程・人月・積算根拠』、法人紹介資料、地域・自治体と関わった実績資料",
]


def get_comment_body() -> str:
    raw = subprocess.check_output(
        ["gh", "api", COMMENT_ENDPOINT], text=True, encoding="utf-8"
    )
    return json.loads(raw)["body"]


def between(text: str, start: str, end: str | None) -> str:
    start_index = text.index(start) + len(start)
    end_index = text.index(end, start_index) if end else len(text)
    result = text[start_index:end_index].strip()
    return re.sub(r"\n---\s*$", "", result).strip()


def plan_sections(body: str) -> dict[int, str]:
    form2 = between(body, "\n# 事業提案計画書", "\n# 概算事業費調書")
    markers = {
        1: "## 1　提案事業の名称",
        2: "## 2　提案事業の内容",
        3: "## 3　提案事業の実施による公益的な効果",
        4: "## 4　対象者、人数、規模等",
        5: "## 5　想定スケジュールと内容",
        6: "## 6　必要最低補助額で事業を実施する場合の概要",
        7: "## 7　備考",
    }
    result = {}
    for number in range(1, 8):
        end = markers.get(number + 1)
        result[number] = between(form2, markers[number], end)
    return result


def set_run_font(run, size=9.0, bold=None):
    run.font.name = "ＭＳ 明朝"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "ＭＳ 明朝")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_inline_text(paragraph, text: str, size=9.0):
    text = text.replace("<br>", "\n")
    link_pattern = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
    cursor = 0
    for match in link_pattern.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]), size)
        label, url = match.groups()
        set_run_font(paragraph.add_run(f"{label}（{url}）"), size)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size)


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def set_cell_text(cell, text: str, size=8.5, bold=False):
    clear_cell(cell)
    for i, line in enumerate(text.split("\n")):
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, size, bold)


def markdown_table(lines: list[str], start: int):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [x.strip() for x in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", x) for x in cells):
            rows.append(cells)
        index += 1
    return rows, index


def apply_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")
        borders.append(element)
    tbl_pr.append(borders)


CELL_CONTENT_DXA = 5020  # 様式1・様式2とも本文列は5138dxa。セル余白分を引いた実効幅


TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook",
]


def reorder_tblpr(table):
    """tblPr子要素をOOXMLスキーマ順へ並べ替える（順序違反だとWordがtblWを無視する）。"""
    tbl_pr = table._tbl.tblPr
    children = list(tbl_pr)
    def key(el):
        tag = el.tag.split("}")[-1]
        return TBLPR_ORDER.index(tag) if tag in TBLPR_ORDER else len(TBLPR_ORDER)
    for child in children:
        tbl_pr.remove(child)
    for child in sorted(children, key=key):
        tbl_pr.append(child)


def nested_column_percents(ncols):
    if ncols == 2:
        return [28, 72]
    if ncols == 4:
        return [18, 34, 14, 34]
    return [100.0 / ncols] * ncols


def set_nested_table_width(table, total_dxa=CELL_CONTENT_DXA):
    """ネスト表を親セル実幅のdxa固定幅＋明示的tblGridにする。

    前回の反省: pct指定は様式1の本文セルにtcWが無いため潰れた（様式2はtcW=5138dxaがあり偶然成立）。
    """
    table.autofit = False
    ncols = len(table.columns)
    percents = nested_column_percents(ncols)
    widths = [int(total_dxa * p / 100) for p in percents]
    widths[-1] = total_dxa - sum(widths[:-1])
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblW")) + tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_pr.append(tbl_w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    reorder_tblpr(table)
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, width in zip(grid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            for existing in tc_pr.findall(qn("w:tcW")):
                tc_pr.remove(existing)
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            tc_pr.append(tc_w)


def render_markdown_to_cell(cell, markdown: str):
    clear_cell(cell)
    lines = markdown.splitlines()
    index = 0
    pending = []

    def flush():
        nonlocal pending
        if not pending:
            return
        text = " ".join(x.strip() for x in pending).strip()
        pending = []
        if not text:
            return
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        is_label = bool(re.fullmatch(r"【[^】]+】", text))
        add_inline_text(paragraph, text, 8.5)
        if is_label:
            for run in paragraph.runs:
                run.bold = True

    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or stripped == "---":
            flush()
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush()
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_text(paragraph, heading.group(2), 9.0)
            for run in paragraph.runs:
                run.bold = True
            index += 1
            continue

        if stripped.startswith("|"):
            flush()
            rows, index = markdown_table(lines, index)
            if rows:
                nested = cell.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                apply_table_borders(nested)
                for ri, row in enumerate(rows):
                    for ci, value in enumerate(row):
                        set_cell_text(nested.cell(ri, ci), value, 7.5, ri == 0)
                set_nested_table_width(nested)
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            flush()
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(9)
            add_inline_text(paragraph, f"・{bullet.group(1)}", 8.5)
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            flush()
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(9)
            add_inline_text(paragraph, f"{numbered.group(1)}．{numbered.group(2)}", 8.5)
            index += 1
            continue

        pending.append(lines[index])
        index += 1
    flush()
    if len(cell.paragraphs) == 0:
        cell.add_paragraph()


def fill_plan_table(table, sections: dict[int, str]):
    set_cell_text(table.cell(0, 1), PROJECT_NAME, 9.0)
    for number in range(2, 8):
        render_markdown_to_cell(table.cell(number - 1, 1), sections[number])


def fill_income_table(table):
    values = [
        (
            "補助金（南砺市ふるさと納税等の寄附を活用した補助金）",
            "100,000,000",
            "100,000,000",
            "企業版ふるさと納税及びガバメントクラウドファンディングによる寄附額を上限として充当。",
        ),
        (
            "自己資金",
            "0",
            "0",
            "汎用的な研究開発及び補助対象外経費は、本表に含めず当法人が別途負担。",
        ),
    ]
    for row_index, row_values in enumerate(values, 1):
        for col_index, value in enumerate(row_values):
            set_cell_text(table.cell(row_index, col_index), value, 8.0)
    for row_index in (3, 4):
        for col_index in range(4):
            set_cell_text(table.cell(row_index, col_index), "", 8.0)
    totals = ("合　計", "100,000,000", "100,000,000", "")
    for col_index, value in enumerate(totals):
        set_cell_text(table.cell(5, col_index), value, 8.0, True)


def fill_expense_table(table):
    values = [
        (
            "人件費",
            "78,000,000",
            "78,000,000",
            "補助事業へ直接従事する65.0人月・10,400時間（1人月160時間・人月単価1,200千円）。"
            "調査・統括12,000,000円、素材制作（工程1）26,400,000円、海外調査・実測（工程2）8,400,000円、"
            "安全性監査（工程3）12,000,000円、来歴確認（工程4B）6,000,000円、人材育成6,000,000円、"
            "測定・公開・報告7,200,000円。作業別内訳は補足資料5.2。",
        ),
        (
            "業務委託費",
            "6,787,000",
            "6,787,000",
            "多言語翻訳・校閲880,000円、撮影1,560,000円、アクセシビリティ確認2,520,000円、"
            "第三者確認1,071,000円、成果物の最終第三者確認756,000円。",
        ),
        (
            "旅費交通費・会場費",
            "4,000,000",
            "4,000,000",
            "長距離出張・市内移動・機材運搬3,600,000円、講座等会場費400,000円。",
        ),
        (
            "海外向け少額実測・システム費・教材等",
            "4,213,000",
            "4,213,000",
            "少額広告2,400,000円、調査参加者謝礼650,000円、教材・印刷・記録媒体・保管990,000円、"
            "AI・GPU・保存・計測配信82,000円、予備・単価変動枠91,000円（未使用分は請求しない）。",
        ),
        (
            "事務・管理費",
            "7,000,000",
            "7,000,000",
            "経理、労務、契約、調達、検収、証憑整理及び実績報告（933時間×7,500円）。"
            "費目・工程別の単価・数量は補足資料第6章。",
        ),
    ]
    for row_index, row_values in enumerate(values, 1):
        for col_index, value in enumerate(row_values):
            set_cell_text(table.cell(row_index, col_index), value, 7.8)
    totals = ("合　計", "100,000,000", "100,000,000", "")
    for col_index, value in enumerate(totals):
        set_cell_text(table.cell(6, col_index), value, 8.0, True)


def replace_project_name_paragraphs(document):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("提案事業の名称"):
            prefix = "提案事業の名称　"
            paragraph.clear()
            set_run_font(paragraph.add_run(prefix + PROJECT_NAME), 9.0)


def fill_form01(source: Path, output: Path, sections: dict[int, str]):
    document = Document(source)
    main = document.tables[0]
    set_cell_text(main.cell(0, 1), PROJECT_NAME, 9.0)
    set_cell_text(main.cell(4, 1), "令和9年4月から令和10年3月まで", 9.0)
    set_cell_text(main.cell(5, 2), "100,000 千円", 9.0)
    clear_cell(main.cell(6, 2))
    amount = main.cell(6, 2).add_paragraph()
    set_run_font(amount.add_run("5,000 千円"), 9.0, True)
    note = main.cell(6, 2).add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    set_run_font(note.add_run(MINIMUM_NOTE), 7.5)
    set_cell_text(main.cell(7, 1), "\n".join(ATTACHMENT_LINES), 8.5)
    fill_plan_table(document.tables[1], sections)
    replace_project_name_paragraphs(document)
    fill_income_table(document.tables[2])
    fill_expense_table(document.tables[3])
    document.save(output)


def fill_form02(source: Path, output: Path, sections: dict[int, str]):
    document = Document(source)
    fill_plan_table(document.tables[0], sections)
    document.save(output)


def fill_form03(source: Path, output: Path):
    document = Document(source)
    replace_project_name_paragraphs(document)
    fill_income_table(document.tables[0])
    fill_expense_table(document.tables[1])
    document.save(output)


def fill_form04(source: Path, output: Path):
    shutil.copy2(source, output)


def document_text(path: Path) -> str:
    document = Document(path)
    chunks = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def nested_tables_full_width(path: Path) -> tuple[int, int]:
    """(ネスト表総数, フル幅設定済み数)"""
    document = Document(path)
    total = full = 0
    for table in document.tables:
        for nested in table._tbl.findall(".//" + qn("w:tbl")):
            total += 1
            tbl_w = nested.find(qn("w:tblPr") + "/" + qn("w:tblW"))
            if tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa" and int(tbl_w.get(qn("w:w"))) >= 5000:
                full += 1
    return total, full


def validate(original: Path, output: Path, required: list[str]):
    original_doc = Document(original)
    output_doc = Document(output)
    if len(original_doc.sections) != len(output_doc.sections):
        raise RuntimeError(f"{output.name}: section count changed")
    if len(original_doc.tables) != len(output_doc.tables):
        raise RuntimeError(f"{output.name}: top-level table count changed")
    for index, (old_table, new_table) in enumerate(zip(original_doc.tables, output_doc.tables)):
        if len(old_table.rows) != len(new_table.rows) or len(old_table.columns) != len(new_table.columns):
            raise RuntimeError(f"{output.name}: table {index} dimensions changed")
    text = document_text(output)
    missing = [term for term in required if term not in text]
    if missing:
        raise RuntimeError(f"{output.name}: missing {missing}")
    total, full = nested_tables_full_width(output)
    if total != full:
        raise RuntimeError(f"{output.name}: nested tables not full width ({full}/{total})")
    return {
        "file": str(output),
        "bytes": output.stat().st_size,
        "tables": len(output_doc.tables),
        "nested_tables_full_width": f"{full}/{total}",
    }


def main():
    body = get_comment_body()
    sections = plan_sections(body)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    fill_form01(*FILES["01"], sections)
    fill_form02(*FILES["02"], sections)
    fill_form03(*FILES["03"])
    fill_form04(*FILES["04"])

    results = [
        validate(
            *FILES["01"],
            [
                PROJECT_NAME,
                "5,000 千円",
                "補助対象経費の下限（5,000千円）に合わせて設定",
                "補足資料『補助額別実施内容・工程・人月・積算根拠』",
                "☑ガバメントクラウドファンディング活用型",
                "100,000,000",
                "Visually",
            ],
        ),
        validate(
            *FILES["02"],
            [PROJECT_NAME, "Visually", "Eris", "SPECA", "Atlas Prover", "補足資料3.3", "補足資料5.2"],
        ),
        validate(
            *FILES["03"],
            [PROJECT_NAME, "78,000,000", "65.0人月・10,400時間", "100,000,000", "6,787,000", "予備・単価変動枠91,000円"],
        ),
        validate(
            *FILES["04"],
            ["一般社団法人 Nyx Foundation", "代表理事", "鳥越 一平"],
        ),
    ]
    print(json.dumps({"output_dir": str(OUTPUT_DIR), "files": results}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
